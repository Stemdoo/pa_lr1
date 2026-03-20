from fuzzywuzzy import fuzz
import docx
import time 
import logging
from Levenshtein import distance as dist

logging.basicConfig(filename = 'res.log', filemode = 'w', level = logging.INFO, encoding="utf-8")

f1 = docx.Document('f1.docx')
f2 = docx.Document('f2.docx')

for ph in f1.paragraphs:
    f1 = ph.text 

for ph in f2.paragraphs:
    f2 = ph.text 

def leven(str_1, str_2):
    n, m = len(str_1), len(str_2)
    if n > m:
        str_1, str_2 = str_2, str_1
        n, m = m, n

    current_row = range(n + 1)
    for i in range(1, m + 1):
        previous_row, current_row = current_row, [i] + [0] * n
        for j in range(1, n + 1):
            add, delete, change = previous_row[j] + 1, current_row[j - 1] + 1, previous_row[j - 1]
            if str_1[j - 1] != str_2[i - 1]:
                change += 1
            current_row[j] = min(add, delete, change)

    return current_row[n]

start = time.perf_counter()
a = fuzz.ratio(f1, f2)
print(a)
end = time.perf_counter()
logging.info(f'функция fuzz.ratio выполнена за : {end - start}')

start = time.perf_counter()
a = dist(f1, f2)
print(a)
end = time.perf_counter()
tm = end - start
logging.info(f'python-Levenshtein: {tm}')

start = time.perf_counter()
a = leven(f1,  f2)
print(a)
end = time.perf_counter()
logging.info(f'моя функция выполнена за :{end - start}')